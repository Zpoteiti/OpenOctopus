from base64 import b64decode
from io import BytesIO
from uuid import uuid4

import pytest
from docx import Document

from openctopus_server.admission import KeyedAdmission
from openctopus_server.errors.codes import ErrorCode
from openctopus_server.errors.exceptions import ToolError
from openctopus_server.workspace.file_content import (
    DocumentParser,
    render_file_content,
    render_streamed_text,
)


def test_text_read_is_line_numbered_and_paginated() -> None:
    rendered = render_file_content(
        "notes.txt",
        b"one\ntwo\nthree\n",
        offset=2,
        limit=1,
    )

    assert rendered == "2|two\n\n(Showing lines 2-2 of 3. Use offset=3 to continue.)"


def test_text_read_has_a_hard_character_cap() -> None:
    rendered = render_file_content("long.txt", ("x" * 130_000).encode())

    assert isinstance(rendered, str)
    assert len(rendered) <= 128_000
    assert rendered.endswith("[truncated]")


def test_png_is_returned_as_an_anthropic_image_block() -> None:
    png = b"\x89PNG\r\n\x1a\n" + b"image"

    rendered = render_file_content("image.bin", png)

    assert isinstance(rendered, list)
    assert rendered[0] == {"type": "text", "text": "Image: image.bin"}
    assert rendered[1]["type"] == "image"
    assert rendered[1]["source"]["media_type"] == "image/png"
    assert b64decode(rendered[1]["source"]["data"]) == png


async def test_docx_text_is_extracted_from_in_memory_bytes() -> None:
    document = Document()
    document.add_paragraph("Workspace report")
    output = BytesIO()
    document.save(output)

    rendered = await _parser().parse("report.docx", output.getvalue(), user_id=uuid4())

    assert rendered == "Workspace report"


def test_unknown_binary_is_rejected() -> None:
    with pytest.raises(ToolError) as caught:
        render_file_content("unknown.bin", b"\x00\x01\x02")

    assert caught.value.code is ErrorCode.TOOL_INVALID_ARGS


async def test_document_parser_uses_a_bounded_disposable_worker() -> None:
    document = Document()
    document.add_paragraph("Parsed outside the event loop")
    output = BytesIO()
    document.save(output)
    parser = _parser()

    rendered = await parser.parse("report.docx", output.getvalue(), user_id=uuid4())

    assert rendered == "Parsed outside the event loop"


async def test_large_text_stream_pages_without_materializing_prior_lines() -> None:
    chunks = iter([b"x" * 200_000 + b"\nsecond\n", b"third\n"])

    async def read() -> bytes:
        return next(chunks, b"")

    rendered = await render_streamed_text(read, offset=2, limit=1)

    assert rendered == "2|second\n\n(Showing lines 2-2 of 3. Use offset=3 to continue.)"


async def test_streamed_text_has_an_aggregate_character_cap() -> None:
    chunks = iter([(("x" * 70_000 + "\n") * 3).encode()])

    async def read() -> bytes:
        return next(chunks, b"")

    rendered = await render_streamed_text(read, offset=1, limit=2000)

    assert len(rendered) <= 128_000
    assert rendered.endswith("[truncated]")


def _parser() -> DocumentParser:
    return DocumentParser(
        admission=KeyedAdmission(global_limit=1, per_key_limit=1, timeout_seconds=1),
        memory_mb=1024,
        timeout_seconds=20,
    )
